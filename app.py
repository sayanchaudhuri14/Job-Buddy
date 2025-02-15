import os
import streamlit as st
from datetime import datetime
from utils.document_parser import extract_candidate_info
from resume_generator.resume_chain import generate_resume
from resume_generator.cover_letter_chain import generate_cover_letter
from docx import Document
import docx
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Page configuration - MUST BE THE FIRST STREAMLIT COMMAND
st.set_page_config(
    page_title="Job Buddy",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
    <style>
    .stTextInput, .stTextArea {
        padding: 10px;
        border-radius: 5px;
        min-height: 40px;
    }
    .stButton button {
        width: 100%;
        border-radius: 5px;
        height: 45px;
    }
    .main .block-container {
        padding-top: 2rem;
    }
    .error-msg {
        color: red;
        font-weight: bold;
    }
    </style>
""", unsafe_allow_html=True)

# Initialize session states
if 'show_form' not in st.session_state:
    st.session_state.show_form = False
if 'experience_entries' not in st.session_state:
    st.session_state.experience_entries = []
if 'education_entries' not in st.session_state:
    st.session_state.education_entries = []
if 'skills_entries' not in st.session_state:
    st.session_state.skills_entries = []
if 'projects_entries' not in st.session_state:
    st.session_state.projects_entries = []
if 'patents_entries' not in st.session_state:
    st.session_state.patents_entries = []
if 'publications_entries' not in st.session_state:
    st.session_state.publications_entries = []
if 'courses_entries' not in st.session_state:
    st.session_state.courses_entries = []
if 'generated_resume' not in st.session_state:
    st.session_state.generated_resume = None
if 'generated_cover_letter' not in st.session_state:
    st.session_state.generated_cover_letter = None

# Sidebar with mandatory key details
with st.sidebar:
    st.title("📋 Key Details")
    st.markdown("---")
    
    # Personal Information
    name = st.text_input("Full Name*", key="name", placeholder="John Doe")
    email = st.text_input("Email Address*", key="email", placeholder="john@example.com")
    contact = st.text_input("Contact Number*", key="contact", placeholder="+1 (123) 456-7890")
    company_name = st.text_input("Target Company*", key="company_name", placeholder="Enter Here")
    st.markdown("*Required fields")
    st.markdown("---")

    st.markdown("### Professional Profiles")
    linkedin = st.text_input("LinkedIn Profile", key="linkedin", placeholder="linkedin.com/in/johndoe")
    github = st.text_input("GitHub Profile", key="github", placeholder="github.com/johndoe")
    other_links = st.text_area("Other Links", key="other_links", 
                              placeholder="Portfolio, Blog, etc.",
                              height=100)
    

    st.markdown("---")
    st.markdown("*Required fields")

# Main content
st.title('🚀 Job Buddy')
st.markdown("Create your professional resume and cover letter tailored to your target job.")

# Key details dictionary
key_details = {
    "name": name,
    "contact": contact,
    "email": email,
    "linkedin": linkedin,
    "github": github,
    "other_links": other_links,
    "company_name": company_name,
}

# Required Fields Check Function
def check_key_details():
    return all([name, email, contact, company_name])
def check_key_details_resume_download():
    return all([name, email, contact])

# Upload Resume Section (Mandatory)
st.markdown("## 📤 Upload Your Resume")
uploaded_file = st.file_uploader(
    "Upload your resume (DOCX or PDF)*",
    type=["docx", "pdf"],
    help="Upload your existing resume to generate a tailored version"
)

# Job Description Section (Mandatory)
st.markdown("## 📋 Job Description")
job_description = st.text_area(
    "Paste the job description*",
    height=200,
    placeholder="Paste the job description here to generate a tailored resume and cover letter..."
)

# LaTeX Template Section
st.markdown("## 📝 LaTeX Template")
template = st.text_area(
    "Paste the LaTeX template code (Required for Resume Generation)*",
    height=200,
    placeholder="Paste your LaTeX template here..."
)

# Enter Information Section
st.markdown("## 📝 Additional Resume Information")
if st.button("Enter Your Information", type="primary", use_container_width=True):
    st.session_state.show_form = not st.session_state.show_form

if st.session_state.show_form:
    # Experience Section
    st.markdown("### Experience")
    if st.button("➕ Add Experience", key="add_exp"):
        st.session_state.experience_entries.append({
            'org_name': '',
            'role': '',
            'start_date': None,
            'end_date': None,
            'description': ''
        })

    for idx, exp in enumerate(st.session_state.experience_entries):
        with st.container():
            st.markdown(f"#### Experience {idx + 1}")
            col1, col2 = st.columns(2)
            with col1:
                exp['org_name'] = st.text_input("Organization Name", key=f"exp_org_{idx}")
                exp['role'] = st.text_input("Role", key=f"exp_role_{idx}")
            with col2:
                exp['start_date'] = st.date_input("Start Date", key=f"exp_start_{idx}")
                exp['end_date'] = st.date_input("End Date", key=f"exp_end_{idx}")
            exp['description'] = st.text_area("Description", key=f"exp_desc_{idx}")
            if st.button("🗑️ Remove", key=f"remove_exp_{idx}"):
                st.session_state.experience_entries.pop(idx)
                st.rerun()

    # Education Section
    st.markdown("### Education")
    if st.button("➕ Add Education", key="add_edu"):
        st.session_state.education_entries.append({
            'school_name': '',
            'degree': '',
            'major': '',
            'start_date': None,
            'end_date': None,
            'courses': []
        })

    for idx, edu in enumerate(st.session_state.education_entries):
        with st.container():
            st.markdown(f"#### Education {idx + 1}")
            col1, col2 = st.columns(2)
            with col1:
                edu['school_name'] = st.text_input("School Name", key=f"edu_school_{idx}")
                edu['degree'] = st.text_input("Degree", key=f"edu_degree_{idx}")
            with col2:
                edu['major'] = st.text_input("Major", key=f"edu_major_{idx}")
                edu['start_date'] = st.date_input("Start Date", key=f"edu_start_{idx}")
                edu['end_date'] = st.date_input("End Date", key=f"edu_end_{idx}")
# Changes for Education Courses Section (around line 200-230)
            # Courses subsection within Education
            st.markdown("##### Courses")
            if st.button("➕ Add Course", key=f"add_course_{idx}"):
                if 'courses' not in edu:
                    edu['courses'] = []
                edu['courses'].append('')
            
            # Courses grid layout
            courses_per_row = 2
            for i in range(0, len(edu.get('courses', [])), courses_per_row):
                cols = st.columns(courses_per_row)
                for j in range(courses_per_row):
                    course_idx = i + j
                    if course_idx < len(edu['courses']):
                        with cols[j]:
                            col1, col2 = st.columns([4, 1])
                            with col1:
                                edu['courses'][course_idx] = st.text_input(
                                    "Course Name",
                                    value=edu['courses'][course_idx],
                                    key=f"course_{idx}_{course_idx}",
                                    placeholder="e.g., Machine Learning"
                                )
                            with col2:
                                if st.button("🗑️", key=f"remove_course_{idx}_{course_idx}"):
                                    edu['courses'].pop(course_idx)
                                    st.rerun()
            if st.button("🗑️ Remove Education", key=f"remove_edu_{idx}"):
                st.session_state.education_entries.pop(idx)
                st.rerun()

    # Changes for Skills Section (around line 250-280 in your code)
    st.markdown("### Skills")
    if st.button("➕ Add Skill Category", use_container_width=True):
        st.session_state.skills_entries.append({
            'category': '',
            'skills': []
        })

    for idx, skill_cat in enumerate(st.session_state.skills_entries):
        with st.container():
            col1, col2 = st.columns([2, 3])
            with col1:
                skill_cat['category'] = st.text_input(
                    "Category Name",
                    key=f"skill_cat_{idx}",
                    placeholder="e.g., Programming Languages"
                )
            with col2:
                if st.button("➕ Add Skill", key=f"add_skill_{idx}"):
                    skill_cat['skills'].append('')
            
            # Skills grid layout
            skills_per_row = 3
            for i in range(0, len(skill_cat['skills']), skills_per_row):
                cols = st.columns(skills_per_row)
                for j in range(skills_per_row):
                    skill_idx = i + j
                    if skill_idx < len(skill_cat['skills']):
                        with cols[j]:
                            col1, col2 = st.columns([4, 1])
                            with col1:
                                skill_cat['skills'][skill_idx] = st.text_input(
                                    "Skill",
                                    value=skill_cat['skills'][skill_idx],
                                    key=f"skill_{idx}_{skill_idx}",
                                    placeholder="e.g., Python"
                                )
                            with col2:
                                if st.button("🗑️", key=f"remove_skill_{idx}_{skill_idx}"):
                                    skill_cat['skills'].pop(skill_idx)
                                    st.rerun()

            if st.button("🗑️ Remove Category", key=f"remove_skill_cat_{idx}"):
                st.session_state.skills_entries.pop(idx)
                st.rerun()
            st.markdown("---")

    # Projects Section
    st.markdown("### Projects")
    if st.button("➕ Add Project"):
        st.session_state.projects_entries.append({
            'name': '',
            'description': '',
            'technologies': '',
            'link': ''
        })

    for idx, project in enumerate(st.session_state.projects_entries):
        with st.container():
            st.markdown(f"#### Project {idx + 1}")
            project['name'] = st.text_input("Project Name", key=f"proj_name_{idx}")
            project['description'] = st.text_area("Description", key=f"proj_desc_{idx}")
            project['technologies'] = st.text_input("Technologies Used", key=f"proj_tech_{idx}")
            project['link'] = st.text_input("Project Link", key=f"proj_link_{idx}")
            if st.button("🗑️ Remove Project", key=f"remove_proj_{idx}"):
                st.session_state.projects_entries.pop(idx)
                st.rerun()

# Changes for Patents Section (around line 300-320)
    st.markdown("### Patents")
    if st.button("➕ Add Patent", use_container_width=True):
        st.session_state.patents_entries.append({
            'title': '',
            'number': '',
            'date': None,
            'description': ''
        })

    for idx, patent in enumerate(st.session_state.patents_entries):
        with st.container():
            st.markdown(f"#### Patent {idx + 1}")
            col1, col2 = st.columns([2, 1])
            with col1:
                patent['title'] = st.text_input(
                    "Patent Title",
                    key=f"patent_title_{idx}",
                    placeholder="Enter patent title"
                )
            with col2:
                patent['number'] = st.text_input(
                    "Patent Number",
                    key=f"patent_num_{idx}",
                    placeholder="e.g., US123456"
                )
            
            col1, col2 = st.columns([1, 2])
            with col1:
                patent['date'] = st.date_input("Filing Date", key=f"patent_date_{idx}")
            with col2:
                patent['description'] = st.text_area(
                    "Description",
                    key=f"patent_desc_{idx}",
                    placeholder="Brief description of the patent",
                    height=100
                )
            
            if st.button("🗑️ Remove Patent", key=f"remove_patent_{idx}, use_container_width=True)"):
                st.session_state.patents_entries.pop(idx)
                st.rerun()
            st.markdown("---")

# Changes for Publications Section (around line 330-350)
    st.markdown("### Publications")
    if st.button("➕ Add Publication", use_container_width=True):
        st.session_state.publications_entries.append({
            'title': '',
            'authors': '',
            'journal': '',
            'year': '',
            'link': ''
        })

    for idx, pub in enumerate(st.session_state.publications_entries):
        with st.container():
            st.markdown(f"#### Publication {idx + 1}")
            pub['title'] = st.text_input(
                "Publication Title",
                key=f"pub_title_{idx}",
                placeholder="Enter publication title"
            )
            
            col1, col2 = st.columns([3, 1])
            with col1:
                pub['authors'] = st.text_input(
                    "Authors",
                    key=f"pub_authors_{idx}",
                    placeholder="e.g., Smith, J., Johnson, M."
                )
            with col2:
                pub['year'] = st.text_input(
                    "Year",
                    key=f"pub_year_{idx}",
                    placeholder="e.g., 2023"
                )
            
            col1, col2 = st.columns(2)
            with col1:
                pub['journal'] = st.text_input(
                    "Journal/Conference",
                    key=f"pub_journal_{idx}",
                    placeholder="e.g., IEEE Transactions"
                )
            with col2:
                pub['link'] = st.text_input(
                    "Link",
                    key=f"pub_link_{idx}",
                    placeholder="DOI or URL"
                )
            
            if st.button("🗑️ Remove Publication", key=f"remove_pub_{idx}, use_container_width=True)"):
                st.session_state.publications_entries.pop(idx)
                st.rerun()
            st.markdown("---")



def create_docx(key_details, experience_entries, education_entries, skills_entries, projects_entries, patents_entries, publications_entries, courses_entries):
    doc = docx.Document()

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run(key_details['name'])
    title_run.bold = True
    title_run.font.size = docx.shared.Pt(24)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Information
    contact_info = doc.add_paragraph()
    contact_info_run = contact_info.add_run(f"{key_details['email']} | {key_details['contact']}")
    contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Professional Profiles
    if key_details['linkedin'] or key_details['github'] or key_details['other_links']:
        profiles = doc.add_paragraph()
        if key_details['linkedin']:
            profiles.add_run(f"LinkedIn: {key_details['linkedin']} | ")
        if key_details['github']:
            profiles.add_run(f"GitHub: {key_details['github']} | ")
        if key_details['other_links']:
            profiles.add_run(f"Other Links: {key_details['other_links']}")
        profiles.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # Experience
    if experience_entries:
        doc.add_heading("Experience", level=2)
        for exp in experience_entries:
            doc.add_paragraph(f"{exp['role']} at {exp['org_name']}")
            doc.add_paragraph(f"{exp['start_date'].strftime('%B %Y')} - {exp['end_date'].strftime('%B %Y') if exp['end_date'] else 'Present'}")
            doc.add_paragraph(exp['description'])

    # Education
    if education_entries:
        doc.add_heading("Education", level=2)
        for edu in education_entries:
            doc.add_paragraph(f"{edu['degree']} in {edu['major']} from {edu['school_name']}")
            doc.add_paragraph(f"{edu['start_date'].strftime('%B %Y')} - {edu['end_date'].strftime('%B %Y') if edu['end_date'] else 'Present'}")
            if 'courses' in edu and edu['courses']:
                doc.add_paragraph("Relevant Coursework: " + ", ".join(edu['courses']))

    # Skills
    if skills_entries:
        doc.add_heading("Skills", level=2)
        for skill_cat in skills_entries:
            doc.add_paragraph(skill_cat['category'], style='List Bullet')  # Use bullet points for categories
            for skill in skill_cat['skills']:
                doc.add_paragraph(skill, style='List Bullet')  # Use bullet points for skills

    # Projects
    if projects_entries:
        doc.add_heading("Projects", level=2)
        for project in projects_entries:
            doc.add_paragraph(project['name'])
            doc.add_paragraph(project['description'])
            doc.add_paragraph(f"Technologies Used: {project['technologies']}")
            if project['link']:
                doc.add_paragraph(f"Link: {project['link']}")

    # Patents
    if patents_entries:
        doc.add_heading("Patents", level=2)
        for patent in patents_entries:
            doc.add_paragraph(patent['title'])
            doc.add_paragraph(f"Patent Number: {patent['number']}")
            doc.add_paragraph(f"Filing Date: {patent['date'].strftime('%B %Y')}")
            doc.add_paragraph(patent['description'])

    # Publications
    if publications_entries:
        doc.add_heading("Publications", level=2)
        for pub in publications_entries:
            doc.add_paragraph(pub['title'])
            doc.add_paragraph(f"Authors: {pub['authors']}")
            doc.add_paragraph(f"Journal/Conference: {pub['journal']} ({pub['year']})")
            if pub['link']:
                doc.add_paragraph(f"Link: {pub['link']}")

    # Save to BytesIO object
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file

# Download Button
if st.button("Generate Information Document", use_container_width=True):
    if check_key_details_resume_download():  # Check mandatory fields
        docx_file = create_docx(key_details, st.session_state.experience_entries, st.session_state.education_entries, st.session_state.skills_entries, st.session_state.projects_entries, st.session_state.patents_entries, st.session_state.publications_entries, st.session_state.courses_entries)

        # Create the filename
        file_name = f"{key_details['name'].lower().replace(' ', '_')}_information.docx"  # Lowercase, replace spaces with underscores

        st.download_button(
            label="Download Resume",
            data=docx_file,
            file_name=file_name,  # Use the dynamically created filename
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.error("Please fill all mandatory fields (Key Details) and add information using the buttons below before downloading.")

key_details_str = f'''
key_details = {{
    "name": "{name}",
    "contact": "{contact}",
    "email": "{email}",
    "linkedin": "{linkedin}",
    "github": "{github}",
    "other_links": "{other_links}",
    "company_name": "{company_name}"
}}
'''


if not check_key_details():
    st.error("Please fill in all required fields in the sidebar.")
elif not uploaded_file:
    st.error("Please upload your resume.")
elif not job_description:
    st.error("Please paste the job description.")
else:
    col1, col2 = st.columns(2)
    
    with col1:
        if template:  # Only show resume generation if template is provided
            if st.button("🎯 Generate Resume", type="primary", use_container_width=True):
                candidate_info = extract_candidate_info(uploaded_file)
                st.session_state.generated_resume = generate_resume(
                    job_description=job_description,
                    candidate_info=candidate_info,
                    key_details=key_details_str,
                    template=template
                )
            
            if st.session_state.generated_resume:
                # Display the generated resume in a text area
                resume_output = st.text_area(
                    "Generated Resume (LaTeX)",
                    value=st.session_state.generated_resume,
                    height=300
                )
                
        else:
            st.warning("Please provide a LaTeX template to generate the resume.")

    with col2:
        if st.button("✉️ Generate Cover Letter", type="primary", use_container_width=True):
            candidate_info = extract_candidate_info(uploaded_file)
            st.session_state.generated_cover_letter = generate_cover_letter(
                job_description=job_description,
                candidate_info=candidate_info,
                key_details=key_details_str
            )
        
        if st.session_state.generated_cover_letter:
            # Display the generated cover letter in a text area
            cover_letter_output = st.text_area(
                "Generated Cover Letter",
                value=st.session_state.generated_cover_letter,
                height=300
            )
           